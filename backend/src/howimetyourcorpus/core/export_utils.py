"""Export corpus et résultats KWIC (TXT, CSV, JSON, JSONL segmenté, Word)."""

from __future__ import annotations

import csv
import json
from pathlib import Path

# Caractères déclenchant des formules dans Excel/LibreOffice si en tête de cellule
_CSV_FORMULA_CHARS = frozenset("=+-@\t\r")


def _csv_safe(value: object) -> object:
    """Protège une valeur CSV contre l'injection de formules (Excel / LibreOffice).

    Les cellules commençant par =, +, -, @, tabulation ou retour chariot
    sont préfixées d'une apostrophe pour être interprétées comme du texte.
    Les valeurs non-string (int, float, None…) sont retournées telles quelles :
    elles ne peuvent pas contenir de formule.
    """
    if not isinstance(value, str):
        return value
    if value and value[0] in _CSV_FORMULA_CHARS:
        return f"'{value}"
    return value

from docx import Document

from howimetyourcorpus.core.models import EpisodeRef
from howimetyourcorpus.core.segment import segmenter_sentences, segmenter_utterances
from howimetyourcorpus.core.storage.db import KwicHit


def _corpus_row(ref: EpisodeRef, clean_text: str) -> dict:
    return {
        "episode_id": ref.episode_id,
        "season": ref.season,
        "episode": ref.episode,
        "title": ref.title or "",
        "clean_text": clean_text,
    }


def export_corpus_txt(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en TXT : une section par épisode (## id - titre + texte)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            f.write(f"## {ref.episode_id} - {ref.title or ''}\n\n")
            f.write(text.strip())
            f.write("\n\n")
    return None


def export_corpus_csv(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en CSV : episode_id, season, episode, title, clean_text."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "clean_text"])
        for ref, text in episodes:
            w.writerow([
                _csv_safe(ref.episode_id),
                ref.season,
                ref.episode,
                _csv_safe(ref.title or ""),
                _csv_safe(text),
            ])
    return None


def export_corpus_json(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en JSON : liste d'objets { episode_id, season, episode, title, clean_text }."""
    data = [_corpus_row(ref, text) for ref, text in episodes]
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return None


def export_corpus_docx(episodes: list[tuple[EpisodeRef, str]], path: Path) -> None:
    """Exporte le corpus en Word (.docx) : un titre par épisode, puis le texte en paragraphes."""
    doc = Document()
    doc.add_heading("Corpus exporté", 0)
    for ref, text in episodes:
        doc.add_heading(f"{ref.episode_id} — {ref.title or ''}", level=1)
        for block in (t.strip() for t in text.split("\n\n") if t.strip()):
            doc.add_paragraph(block)
        doc.add_paragraph()
    doc.save(str(path))
    return None


# ----- Export segments (depuis l'Inspecteur) -----

SEGMENT_EXPORT_COLUMNS = ["segment_id", "episode_id", "kind", "n", "start_char", "end_char", "text"]


def export_segments_txt(segments: list[dict], path: Path) -> None:
    """Exporte les segments en TXT : un segment par ligne (texte uniquement)."""
    with path.open("w", encoding="utf-8") as f:
        for s in segments:
            text = (s.get("text") or "").strip().replace("\n", " ")
            f.write(text)
            f.write("\n")
    return None


def export_segments_csv(segments: list[dict], path: Path) -> None:
    """Exporte les segments en CSV : segment_id, episode_id, kind, n, start_char, end_char, text."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(SEGMENT_EXPORT_COLUMNS)
        for s in segments:
            w.writerow([
                _csv_safe(s.get("segment_id", "")),
                _csv_safe(s.get("episode_id", "")),
                _csv_safe(s.get("kind", "")),
                s.get("n", ""),
                s.get("start_char", ""),
                s.get("end_char", ""),
                _csv_safe((s.get("text") or "").replace("\n", " ")),
            ])
    return None


def export_segments_tsv(segments: list[dict], path: Path) -> None:
    """Exporte les segments en TSV : même colonnes que CSV."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        w.writerow(SEGMENT_EXPORT_COLUMNS)
        for s in segments:
            w.writerow([
                _csv_safe(s.get("segment_id", "")),
                _csv_safe(s.get("episode_id", "")),
                _csv_safe(s.get("kind", "")),
                s.get("n", ""),
                s.get("start_char", ""),
                s.get("end_char", ""),
                _csv_safe((s.get("text") or "").replace("\n", " ")),
            ])
    return None


def export_segments_docx(segments: list[dict], path: Path) -> None:
    """Exporte les segments en Word (.docx) : tableau segment_id, episode_id, kind, n, text."""
    doc = Document()
    doc.add_heading("Segments exportés", 0)
    if not segments:
        doc.add_paragraph("Aucun segment.")
        doc.save(str(path))
        return None
    table = doc.add_table(rows=1 + len(segments), cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "segment_id"
    h[1].text = "episode_id"
    h[2].text = "kind"
    h[3].text = "text"
    for i, s in enumerate(segments):
        row = table.rows[i + 1].cells
        row[0].text = str(s.get("segment_id", ""))
        row[1].text = str(s.get("episode_id", ""))
        row[2].text = str(s.get("kind", ""))
        row[3].text = (s.get("text") or "").replace("\n", " ").strip()
    doc.save(str(path))
    return None


def export_segments_srt_like(segments: list[dict], path: Path) -> None:
    """Exporte les segments en format SRT-like : blocs numérotés (timecodes 00:00:00,000 si absents)."""
    with path.open("w", encoding="utf-8") as f:
        for i, s in enumerate(segments, start=1):
            text = (s.get("text") or "").strip().replace("\n", " ")
            f.write(f"{i}\n")
            f.write("00:00:00,000 --> 00:00:00,000\n")
            f.write(text)
            f.write("\n\n")
    return None


def export_kwic_csv(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en CSV (inclut segment_id/kind/cue_id/lang/speaker si présents)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        row0 = ["episode_id", "title", "left", "match", "right", "position", "score"]
        if hits and (getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None) or getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None) or getattr(hits[0], "speaker", None)):
            if getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None):
                row0.extend(["segment_id", "kind"])
            if getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None):
                row0.extend(["cue_id", "lang"])
            if getattr(hits[0], "speaker", None):
                row0.append("speaker")
        w.writerow(row0)
        for h in hits:
            r = [_csv_safe(h.episode_id), _csv_safe(h.title), _csv_safe(h.left), _csv_safe(h.match), _csv_safe(h.right), h.position, h.score]
            if len(row0) > 7:
                r.extend([_csv_safe(getattr(h, "segment_id", "") or ""), _csv_safe(getattr(h, "kind", "") or "")])
            if len(row0) > 9:
                r.extend([_csv_safe(getattr(h, "cue_id", "") or ""), _csv_safe(getattr(h, "lang", "") or "")])
            if "speaker" in row0:
                r.append(_csv_safe(getattr(h, "speaker", "") or ""))
            w.writerow(r)
    return None


def export_kwic_tsv(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en TSV (inclut segment_id/kind/cue_id/lang/speaker si présents)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        row0 = ["episode_id", "title", "left", "match", "right", "position", "score"]
        if hits and (getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None) or getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None) or getattr(hits[0], "speaker", None)):
            if getattr(hits[0], "segment_id", None) or getattr(hits[0], "kind", None):
                row0.extend(["segment_id", "kind"])
            if getattr(hits[0], "cue_id", None) or getattr(hits[0], "lang", None):
                row0.extend(["cue_id", "lang"])
            if getattr(hits[0], "speaker", None):
                row0.append("speaker")
        w.writerow(row0)
        for h in hits:
            r = [_csv_safe(h.episode_id), _csv_safe(h.title), _csv_safe(h.left), _csv_safe(h.match), _csv_safe(h.right), h.position, h.score]
            if len(row0) > 7:
                r.extend([_csv_safe(getattr(h, "segment_id", "") or ""), _csv_safe(getattr(h, "kind", "") or "")])
            if len(row0) > 9:
                r.extend([_csv_safe(getattr(h, "cue_id", "") or ""), _csv_safe(getattr(h, "lang", "") or "")])
            if "speaker" in row0:
                r.append(_csv_safe(getattr(h, "speaker", "") or ""))
            w.writerow(r)
    return None


def export_kwic_json(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en JSON (inclut segment_id/kind si présents)."""
    data = []
    for h in hits:
        row = {
            "episode_id": h.episode_id,
            "title": h.title,
            "left": h.left,
            "match": h.match,
            "right": h.right,
            "position": h.position,
            "score": h.score,
        }
        if getattr(h, "segment_id", None):
            row["segment_id"] = h.segment_id
        if getattr(h, "kind", None):
            row["kind"] = h.kind
        if getattr(h, "cue_id", None):
            row["cue_id"] = h.cue_id
        if getattr(h, "lang", None):
            row["lang"] = h.lang
        if getattr(h, "speaker", None):
            row["speaker"] = h.speaker
        data.append(row)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return None


def export_kwic_jsonl(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en JSONL (une ligne JSON par hit)."""
    with path.open("w", encoding="utf-8") as f:
        for h in hits:
            row = {
                "episode_id": h.episode_id,
                "title": h.title,
                "left": h.left,
                "match": h.match,
                "right": h.right,
                "position": h.position,
                "score": h.score,
            }
            if getattr(h, "segment_id", None):
                row["segment_id"] = h.segment_id
            if getattr(h, "kind", None):
                row["kind"] = h.kind
            if getattr(h, "cue_id", None):
                row["cue_id"] = h.cue_id
            if getattr(h, "lang", None):
                row["lang"] = h.lang
            if getattr(h, "speaker", None):
                row["speaker"] = h.speaker
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    return None


def _kwic_row_values(h: KwicHit, cols: list[str]) -> list[str]:
    """Retourne les valeurs d'un hit KWIC dans l'ordre des colonnes."""
    values = []
    for c in cols:
        if c == "episode_id":
            values.append(h.episode_id or "")
        elif c == "title":
            values.append(h.title or "")
        elif c == "left":
            values.append((h.left or "").replace("\n", " "))
        elif c == "match":
            values.append((h.match or "").replace("\n", " "))
        elif c == "right":
            values.append((h.right or "").replace("\n", " "))
        elif c == "position":
            values.append(str(h.position))
        elif c == "score":
            values.append(str(h.score))
        elif c == "segment_id":
            values.append(getattr(h, "segment_id", "") or "")
        elif c == "kind":
            values.append(getattr(h, "kind", "") or "")
        elif c == "cue_id":
            values.append(getattr(h, "cue_id", "") or "")
        elif c == "lang":
            values.append(getattr(h, "lang", "") or "")
        elif c == "speaker":
            values.append(getattr(h, "speaker", "") or "")
        else:
            values.append("")
    return values


def export_kwic_docx(hits: list[KwicHit], path: Path) -> None:
    """Exporte les résultats KWIC en Word (.docx) : tableau left, match, right + métadonnées."""
    doc = Document()
    doc.add_heading("Résultats KWIC", 0)
    if not hits:
        doc.add_paragraph("Aucun résultat.")
        doc.save(str(path))
        return None
    h0 = hits[0]
    cols = ["episode_id", "title", "left", "match", "right", "position", "score"]
    if getattr(h0, "segment_id", None) or getattr(h0, "kind", None):
        cols.extend(["segment_id", "kind"])
    if getattr(h0, "cue_id", None) or getattr(h0, "lang", None):
        cols.extend(["cue_id", "lang"])
    if getattr(h0, "speaker", None):
        cols.append("speaker")
    table = doc.add_table(rows=1 + len(hits), cols=len(cols))
    table.style = "Table Grid"
    for j, c in enumerate(cols):
        table.rows[0].cells[j].text = c
    for i, h in enumerate(hits):
        for j, val in enumerate(_kwic_row_values(h, cols)):
            table.rows[i + 1].cells[j].text = val
    doc.save(str(path))
    return None


# --- Export segmenté (Phase 2 : utterances / phrases, JSONL + CSV) ---


def _iter_utterance_rows(clean_text: str) -> list[dict]:
    """Construit des lignes d'utterances à partir du segmenter courant (non legacy)."""
    rows: list[dict] = []
    for seg in segmenter_utterances(clean_text):
        rows.append(
            {
                "speaker": seg.speaker_explicit,
                "text": seg.text,
                "index": seg.n,
            }
        )
    return rows


def _iter_phrase_rows(clean_text: str) -> list[dict]:
    """
    Construit des lignes de phrases à partir du segmenter courant.
    Le speaker est propagé depuis l'utterance d'origine.
    """
    rows: list[dict] = []
    idx = 0
    for utt in segmenter_utterances(clean_text):
        sentences = segmenter_sentences(utt.text, lang_hint="en")
        if not sentences:
            text = (utt.text or "").strip()
            if text:
                rows.append({"speaker": utt.speaker_explicit, "text": text, "index": idx})
                idx += 1
            continue
        for sent in sentences:
            rows.append({"speaker": utt.speaker_explicit, "text": sent.text, "index": idx})
            idx += 1
    return rows


def export_corpus_utterances_jsonl(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en utterances : une ligne JSON par utterance (JSONL)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            for u in _iter_utterance_rows(text):
                obj = {
                    "episode_id": ref.episode_id,
                    "season": ref.season,
                    "episode": ref.episode,
                    "title": ref.title or "",
                    "speaker": u.get("speaker"),
                    "text": u.get("text", ""),
                    "index": u.get("index", 0),
                }
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
    return None


def export_corpus_utterances_csv(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en utterances : CSV (episode_id, season, episode, title, speaker, text, index)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "speaker", "text", "index"])
        for ref, text in episodes:
            for u in _iter_utterance_rows(text):
                w.writerow([
                    _csv_safe(ref.episode_id),
                    ref.season,
                    ref.episode,
                    _csv_safe(ref.title or ""),
                    _csv_safe(u.get("speaker") or ""),
                    _csv_safe(u.get("text", "")),
                    u.get("index", 0),
                ])
    return None


def export_corpus_phrases_jsonl(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en phrases : une ligne JSON par phrase (JSONL)."""
    with path.open("w", encoding="utf-8") as f:
        for ref, text in episodes:
            for ph in _iter_phrase_rows(text):
                obj = {
                    "episode_id": ref.episode_id,
                    "season": ref.season,
                    "episode": ref.episode,
                    "title": ref.title or "",
                    "speaker": ph.get("speaker"),
                    "text": ph.get("text", ""),
                    "index": ph.get("index", 0),
                }
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")
    return None


def export_corpus_phrases_csv(
    episodes: list[tuple[EpisodeRef, str]], path: Path
) -> None:
    """Exporte le corpus segmenté en phrases : CSV (episode_id, season, episode, title, speaker, text, index)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["episode_id", "season", "episode", "title", "speaker", "text", "index"])
        for ref, text in episodes:
            for ph in _iter_phrase_rows(text):
                w.writerow([
                    _csv_safe(ref.episode_id),
                    ref.season,
                    ref.episode,
                    _csv_safe(ref.title or ""),
                    _csv_safe(ph.get("speaker") or ""),
                    _csv_safe(ph.get("text", "")),
                    ph.get("index", 0),
                ])
    return None


# --- Phase 5 : concordancier parallèle et rapports ---

PARALLEL_CONCORDANCE_COLUMNS = [
    "segment_id", "speaker", "text_segment", "text_en", "confidence_pivot",
    "text_fr", "confidence_fr", "text_it", "confidence_it",
]


def _parallel_cell(row: dict, key: str):
    """Valeur d'une cellule pour l'export CSV/TSV : None → chaîne vide, protection injection CSV."""
    v = row.get(key)
    if v is None:
        return ""
    return _csv_safe(v)


def export_parallel_concordance_csv(rows: list[dict], path: Path) -> None:
    """Exporte le concordancier parallèle en CSV (segment, EN, FR + confiances)."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(PARALLEL_CONCORDANCE_COLUMNS)
        for r in rows:
            w.writerow([_parallel_cell(r, k) for k in PARALLEL_CONCORDANCE_COLUMNS])
    return None


def export_parallel_concordance_tsv(rows: list[dict], path: Path) -> None:
    """Exporte le concordancier parallèle en TSV."""
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        w.writerow(PARALLEL_CONCORDANCE_COLUMNS)
        for r in rows:
            w.writerow([_parallel_cell(r, k) for k in PARALLEL_CONCORDANCE_COLUMNS])
    return None


def export_parallel_concordance_jsonl(rows: list[dict], path: Path) -> None:
    """Exporte le concordancier parallèle en JSONL (une ligne JSON par alignement)."""
    with path.open("w", encoding="utf-8") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")
    return None


def export_parallel_concordance_docx(rows: list[dict], path: Path) -> None:
    """Exporte le concordancier parallèle en Word (.docx) : tableau segment, EN, FR, IT + confiances."""
    doc = Document()
    doc.add_heading("Concordancier parallèle", 0)
    if not rows:
        doc.add_paragraph("Aucune ligne.")
        doc.save(str(path))
        return None
    table = doc.add_table(rows=1 + len(rows), cols=len(PARALLEL_CONCORDANCE_COLUMNS))
    table.style = "Table Grid"
    for j, col in enumerate(PARALLEL_CONCORDANCE_COLUMNS):
        table.rows[0].cells[j].text = col
    for i, r in enumerate(rows):
        for j, col in enumerate(PARALLEL_CONCORDANCE_COLUMNS):
            val = r.get(col)
            table.rows[i + 1].cells[j].text = "" if val is None else str(val)
    doc.save(str(path))
    return None


def export_parallel_concordance_txt(rows: list[dict], path: Path) -> None:
    """§15.1 — Exporte la comparaison de traductions en TXT : une ligne par alignement, colonnes séparées par tab (segment | EN | FR | IT)."""
    with path.open("w", encoding="utf-8") as f:
        for r in rows:
            cells = [str(_parallel_cell(r, k)) for k in PARALLEL_CONCORDANCE_COLUMNS]
            f.write("\t".join(cells).replace("\n", " ").replace("\r", ""))
            f.write("\n")
    return None


def export_parallel_concordance_html(rows: list[dict], path: Path, title: str | None = None) -> None:
    """§15.1 — Exporte la comparaison de traductions en HTML : tableau segment | EN | FR | IT (sans stats)."""
    t = title or "Comparaison de traductions"
    lines = [
        "<!DOCTYPE html>",
        "<html><head><meta charset='utf-8'><title>" + _escape(t) + "</title>",
        "<style>table { border-collapse: collapse; } th, td { border: 1px solid #ccc; padding: 6px; text-align: left; }</style>",
        "</head><body>",
        "<h1>" + _escape(t) + "</h1>",
        "<table><thead><tr>",
    ]
    for col in PARALLEL_CONCORDANCE_COLUMNS:
        lines.append("<th>" + _escape(col) + "</th>")
    lines.append("</tr></thead><tbody>")
    for r in rows:
        lines.append("<tr>")
        for col in PARALLEL_CONCORDANCE_COLUMNS:
            val = r.get(col)
            cell = "" if val is None else str(val)
            lines.append("<td>" + _escape(cell) + "</td>")
        lines.append("</tr>")
    lines.append("</tbody></table></body></html>")
    path.write_text("\n".join(lines), encoding="utf-8")
    return None


def export_align_report_html(
    stats: dict,
    sample_rows: list[dict],
    episode_id: str,
    run_id: str,
    path: Path,
    title: str | None = None,
) -> None:
    """
    Génère un rapport HTML (Phase 5) : stats d'alignement + tableau échantillon du concordancier parallèle.
    Quarto reste optionnel pour des rapports avancés.
    """
    t = title or f"Rapport alignement — {episode_id}"
    by_status = stats.get("by_status") or {}
    by_status_pivot = stats.get("by_status_pivot") or {}
    coverage_pct = stats.get("coverage_pct")
    coverage_str = f"{coverage_pct}%" if coverage_pct is not None else "—"
    lines = [
        "<!DOCTYPE html>",
        "<html><head><meta charset='utf-8'><title>" + _escape(t) + "</title></head><body>",
        "<h1>" + _escape(t) + "</h1>",
        "<p><strong>Épisode:</strong> " + _escape(stats.get("episode_id", "")) + "</p>",
        "<p><strong>Run:</strong> " + _escape(stats.get("run_id", "")) + "</p>",
        "<h2>Statistiques</h2>",
        "<ul>",
        "<li>Liens totaux: " + str(stats.get("nb_links", 0)) + "</li>",
        "<li>Liens pivot (segment↔EN): " + str(stats.get("nb_pivot", 0)) + "</li>",
        "<li>Liens target (EN↔FR): " + str(stats.get("nb_target", 0)) + "</li>",
        "<li>Confiance moyenne: " + (str(stats.get("avg_confidence")) if stats.get("avg_confidence") is not None else "—") + "</li>",
        "<li>Couverture (pivot): " + coverage_str + "</li>",
        "<li>Par statut (pivot uniquement): " + ", ".join(f"{k}={v}" for k, v in sorted(by_status_pivot.items())) + "</li>",
        "<li>Par statut (tous rôles): " + ", ".join(f"{k}={v}" for k, v in sorted(by_status.items())) + "</li>",
        "</ul>",
        "<h2>Échantillon concordancier parallèle</h2>",
        "<table border='1' cellpadding='4' style='border-collapse: collapse;'>",
        "<thead><tr><th>segment_id</th><th>Personnage</th><th>Segment (transcript)</th><th>EN</th><th>conf.</th><th>FR</th><th>conf.</th><th>IT</th><th>conf.</th></tr></thead>",
        "<tbody>",
    ]
    for r in sample_rows[:100]:
        t_seg = str(r.get("text_segment", ""))
        t_en = str(r.get("text_en", ""))
        t_fr = str(r.get("text_fr", ""))
        t_it = str(r.get("text_it", ""))
        speaker = str(r.get("speaker", ""))
        lines.append(
            "<tr><td>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td></tr>".format(
                _escape(str(r.get("segment_id", ""))),
                _escape(speaker),
                _escape((t_seg[:80] + "…") if len(t_seg) > 80 else t_seg),
                _escape((t_en[:60] + "…") if len(t_en) > 60 else t_en),
                _escape(str(r.get("confidence_pivot") if r.get("confidence_pivot") is not None else "")),
                _escape((t_fr[:60] + "…") if len(t_fr) > 60 else t_fr),
                _escape(str(r.get("confidence_fr") if r.get("confidence_fr") is not None else "")),
                _escape((t_it[:60] + "…") if len(t_it) > 60 else t_it),
                _escape(str(r.get("confidence_it") if r.get("confidence_it") is not None else "")),
            )
        )
    lines.extend(["</tbody></table>", "</body></html>"])
    path.write_text("\n".join(lines), encoding="utf-8")
    return None


def _escape(s: str | None) -> str:
    """Échappe HTML pour affichage sûr. Accepte None et renvoie ''."""
    if s is None:
        return ""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
